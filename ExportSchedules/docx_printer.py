import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocxPrinter():
    def __init__(self, template, print_path):
        self.print_path = print_path
        self.template = template
        self.NUM_OF_COLUMNS = 7
        self.font_name = 'GOST'
        self.font_size = 12
        print('docx printer initialized')

    def print(self, docx_file, specification = {}):
        doc = Document(docx = self.template)
        table = doc.tables.pop()
        i = 1
        for section, data in specification.items():
            print('printing row:', i, 'section:', section)
            i += 1
            table.rows[i].cells[1].text = section
            run = table.rows[i].cells[1].paragraphs[0].runs[0]
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)
            run.bold = True
            table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            
            for line in data:
                for n in range(self.NUM_OF_COLUMNS):
                    table.rows[i].cells[n].text = line[n]
                    run = table.rows[i].cells[n].paragraphs[0].runs[0]
                    if n > 1 : table.rows[i].cells[n].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.font.name = self.font_name
                    run.font.size = Pt(self.font_size)
                i += 1
                if i == len(table.rows): table.add_row()
                print('\tprinted line', line)

        doc.save(os.path.join(self.print_path, docx_file))
